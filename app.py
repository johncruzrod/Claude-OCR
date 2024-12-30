import streamlit as st
import streamlit.components.v1 as components
import anthropic
import base64
import io
import tempfile
from docx import Document
from PIL import Image

# Configure the page
st.set_page_config(page_title="Image Text Extractor", page_icon="📝", layout="wide")

# -------------------------------------------------------------------
# 1) HTML snippet for camera feed, NO internal capture button
# -------------------------------------------------------------------
# We rely on an external button that sends 'capturePhoto' to this iframe.
CAMERA_HTML = """
<style>
  #camera_wrap {
    display: flex;
    flex-direction: column;
    align-items: center;
    justify-content: center;
  }
  video {
    width: 100%;
    max-width: 600px;
    border: 2px solid #ccc;
    margin-bottom: 1em;
  }
  canvas {
    display: none;
  }
</style>
<div id="camera_wrap">
  <video id="video" playsinline autoplay></video>
  <canvas id="canvas"></canvas>
</div>
<script>
let stream = null;

async function initCamera() {
    const video = document.getElementById('video');
    const canvas = document.getElementById('canvas');

    try {
        // Attempt high-resolution rear camera
        stream = await navigator.mediaDevices.getUserMedia({
            video: {
                facingMode: { ideal: 'environment' },
                width: { ideal: 4096 },
                height: { ideal: 2160 },
                focusMode: { ideal: 'continuous' }
            }
        });
    } catch (err) {
        // Fallback if 4K is blocked
        try {
            stream = await navigator.mediaDevices.getUserMedia({
                video: {
                    width: { ideal: 1920 },
                    height: { ideal: 1080 }
                }
            });
        } catch (fallbackErr) {
            console.error("Camera access failed:", fallbackErr);
            return;
        }
    }

    video.srcObject = stream;
    video.setAttribute("playsinline", true);

    // When the parent sends 'capturePhoto', capture current frame
    window.addEventListener('message', (event) => {
        if (event.data === 'capturePhoto') {
            canvas.width = video.videoWidth;
            canvas.height = video.videoHeight;
            const ctx = canvas.getContext('2d');
            ctx.drawImage(video, 0, 0, canvas.width, canvas.height);
            const imageData = canvas.toDataURL('image/jpeg', 0.95);
            // Return it to parent
            window.parent.postMessage({
                type: 'streamlit:setComponentValue',
                value: imageData
            }, '*');
        }
    });

    // Stop tracks on page unload
    window.addEventListener('beforeunload', () => {
        if (stream) {
            stream.getTracks().forEach(track => track.stop());
        }
    });
}

initCamera();
</script>
"""

# -------------------------------------------------------------------
# 2) Functions for image optimisation and Claude Vision call
# -------------------------------------------------------------------
def optimise_image(image_data, is_base64=False):
    """
    Ensure the image is a JPEG under ~5 MB while preserving as much quality as possible.
    """
    try:
        if is_base64:
            # Strip any data URL prefix
            if "," in image_data:
                image_data = image_data.split(",")[1]
            raw = base64.b64decode(image_data)
            img = Image.open(io.BytesIO(raw))
        else:
            if hasattr(image_data, "seek"):
                image_data.seek(0)
            img = Image.open(image_data)

        # Convert to RGB if necessary
        if img.mode != "RGB":
            img = img.convert("RGB")

        # Resize if over 4096 in any dimension
        max_dim = 4096
        w, h = img.size
        if w > max_dim or h > max_dim:
            if w >= h:
                new_w = max_dim
                new_h = int(h * (max_dim / w))
            else:
                new_h = max_dim
                new_w = int(w * (max_dim / h))
            img = img.resize((new_w, new_h), Image.Resampling.LANCZOS)

        # Save at quality=95, check size
        buf = io.BytesIO()
        img.save(buf, format="JPEG", optimize=True, quality=95)
        size_bytes = buf.tell()

        # If still > 5MB, lower quality in steps
        if size_bytes > 5 * 1024 * 1024:
            quality = 90
            while size_bytes > 5 * 1024 * 1024 and quality >= 75:
                buf = io.BytesIO()
                img.save(buf, format="JPEG", optimize=True, quality=quality)
                size_bytes = buf.tell()
                quality -= 5

        buf.seek(0)
        return buf.getvalue()

    except Exception as e:
        raise Exception(f"Image optimization failed: {str(e)}")

def process_image_with_claude(client, image_bytes):
    """
    Send the image to Claude Vision for text extraction.
    """
    try:
        img_b64 = base64.b64encode(image_bytes).decode()
        prompt = """Please accurately transcribe all text from this image.
Output only the transcribed text with no additional commentary.
Preserve formatting and structure where possible.
Pay special attention to handwriting and use context to ensure accuracy."""

        response = client.messages.create(
            model="claude-3-sonnet-20240229",
            max_tokens=4096,
            temperature=0.0,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": "image/jpeg",
                                "data": img_b64
                            }
                        },
                        {
                            "type": "text",
                            "text": prompt
                        }
                    ]
                }
            ]
        )
        return response.content[0].text
    except Exception as e:
        return f"Error processing image: {str(e)}"

def create_docx(text_list):
    """
    Create a docx containing all extracted texts.
    """
    doc = Document()
    doc.add_heading("Extracted Text from Images", 0)
    for i, text in enumerate(text_list, start=1):
        doc.add_heading(f"Image {i}", level=1)
        doc.add_paragraph(text)
        if i < len(text_list):
            doc.add_page_break()

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return tmp.name

# -------------------------------------------------------------------
# 3) Main App
# -------------------------------------------------------------------
def main():
    st.title("Image Text Extractor")

    # Check we have the Anthropic key
    if "ANTHROPIC_API_KEY" not in st.secrets:
        st.error("Please set st.secrets['ANTHROPIC_API_KEY'] to your Anthropic key.")
        return
    client = anthropic.Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])

    # This is where we store images (each item is dict with {"data": bytes, "method": ...})
    if "images" not in st.session_state:
        st.session_state["images"] = []

    # Choose camera or upload
    method = st.radio("Choose input method:", ["Upload Image(s)", "Take Photo"])

    # --- (A) Upload multiple images ---
    if method == "Upload Image(s)":
        st.info("Upload one or more images for text extraction.")
        files = st.file_uploader("Select images:", accept_multiple_files=True, type=["png", "jpg", "jpeg"])
        if files:
            for f in files:
                # Avoid duplicates
                if f not in [x.get("file") for x in st.session_state["images"] if x.get("file")]:
                    try:
                        optimised = optimise_image(f)
                        st.session_state["images"].append({
                            "data": optimised,
                            "file": f,
                            "method": "upload"
                        })
                    except Exception as e:
                        st.error(f"Error with {f.name}: {str(e)}")

    # --- (B) High-res camera feed with external capture button ---
    else:
        st.info("Ensure good lighting, position the camera, then click 'Capture Photo'.")

        # 1) Insert the camera HTML (no "key" passed)
        camera_component = components.html(CAMERA_HTML, height=600, scrolling=False)

        # 2) A hidden input that is purely for receiving the base64 from the iframe
        st.markdown(
            """
            <input id="hidden_camera_data" type="hidden" />
            <script>
            //  Listen for the "streamlit:setComponentValue" event from the iframe
            window.addEventListener("message", (event) => {
                if(event.data && event.data.type === "streamlit:setComponentValue"){
                    const base64Image = event.data.value;
                    // Put it in our hidden input
                    const hiddenInput = document.getElementById("hidden_camera_data");
                    hiddenInput.value = base64Image;
                    // Trigger a change event so Streamlit sees the updated value
                    hiddenInput.dispatchEvent(new Event('change'));
                }
            });
            </script>
            """,
            unsafe_allow_html=True
        )

        # 3) The external "Capture Photo" button in Streamlit
        if st.button("Capture Photo"):
            # Send a message to the iframe telling it to capture
            st.markdown(
                """
                <script>
                const frames = window.top.document.getElementsByTagName('iframe');
                for(let f of frames){
                    if(f.srcdoc && f.srcdoc.indexOf('initCamera') !== -1){
                        f.contentWindow.postMessage('capturePhoto','*');
                        break;
                    }
                }
                </script>
                """,
                unsafe_allow_html=True
            )

        # 4) Watch for changes to #hidden_camera_data from the script. If changed, get it in Python.
        #    We'll use st.experimental_get_query_params hack or a timer. Easiest is a small callback.
        #    Since Streamlit cannot automatically pick up DOM changes, we do a minimal workaround:
        captured_image_b64 = st.session_state.get("camera_base64", None)

        # We'll poll the hidden input's value with a small "poll" button or an invisible form.
        # Instead, let's do a tiny "Invisible Poll" by running some JavaScript that sets a query param
        # or re-run. A simpler approach is an "on_change" text_input, but user said "No text boxes!"
        # So we'll do a 1-second auto refresh and read the hidden input with `st.js_onclick`.
        # 
        # For reliability, let's just provide a small "Finalize Capture" button. 
        # If you want it completely automatic, you'd do a custom component. 
        if st.button("Finalize Capture"):
            # This snippet reads the hidden input's value from the DOM
            get_js = """
            <script>
            const hiddenValue = document.getElementById("hidden_camera_data").value;
            window.parent.postMessage({
                type: "setSessionStateCamera",
                value: hiddenValue
            }, "*");
            </script>
            """
            st.markdown(get_js, unsafe_allow_html=True)

        # Listen for messages from the parent if we do the above postMessage
        # We'll store in st.session_state["camera_base64"] so we can do the next step
        # We can’t do it purely in Python, so we do a small <script>:
        st.markdown(
            """
            <script>
            window.addEventListener("message", (event) => {
                if(event.data && event.data.type==="setSessionStateCamera"){
                    const base64Str = event.data.value;
                    // Trigger a Streamlit state update via a custom hack:
                    window.location.href = window.location.href.split("?")[0] + "?capturedb64=" + encodeURIComponent(base64Str);
                }
            });
            </script>
            """,
            unsafe_allow_html=True
        )

        # Now we read ?capturedb64= from the URL
        query_params = st.experimental_get_query_params()
        if "capturedb64" in query_params:
            new_b64 = query_params["capturedb64"][0]
            if new_b64.startswith("data:image/jpeg;base64"):
                try:
                    optimised_bytes = optimise_image(new_b64, is_base64=True)
                    # Check duplicates
                    if optimised_bytes not in [x["data"] for x in st.session_state["images"]]:
                        st.session_state["images"].append({"data": optimised_bytes, "method": "camera"})
                    # Clear param from URL by reloading without it
                    st.experimental_set_query_params()
                    st.experimental_rerun()
                except Exception as e:
                    st.error(f"Error processing captured photo: {str(e)}")
            else:
                # Clear param if invalid
                st.experimental_set_query_params()

    # ----------------------------------------------------------------
    # Show images grid
    # ----------------------------------------------------------------
    if st.session_state["images"]:
        st.subheader("Images to Process")
        cols = st.columns(3)
        for i, img_info in enumerate(st.session_state["images"]):
            with cols[i % 3]:
                try:
                    pil_img = Image.open(io.BytesIO(img_info["data"]))
                    st.image(pil_img, caption=f"Image {i + 1}", use_column_width=True)

                    if st.button("Remove", key=f"remove_{i}"):
                        st.session_state["images"].pop(i)
                        st.experimental_rerun()

                except Exception as e:
                    st.error(f"Error displaying image {i+1}: {str(e)}")

        if len(st.session_state["images"]) > 1:
            if st.button("Clear All"):
                st.session_state["images"] = []
                st.experimental_rerun()

        # ----------------------------------------------------------------
        # Extract text with Claude Vision
        # ----------------------------------------------------------------
        if st.button("Extract Text", type="primary"):
            st.subheader("Extracted Text")
            texts = []
            prog = st.progress(0)
            status_area = st.empty()

            for idx, info in enumerate(st.session_state["images"]):
                status_area.write(f"Processing image {idx + 1} of {len(st.session_state['images'])}...")
                result = process_image_with_claude(client, info["data"])
                texts.append(result)

                with st.expander(f"Text from Image {idx + 1}", expanded=True):
                    st.write(result)

                prog.progress((idx + 1) / len(st.session_state["images"]))

            status_area.write("Done!")

            # Download docx
            if texts:
                doc_path = create_docx(texts)
                with open(doc_path, "rb") as f:
                    st.download_button(
                        label="Download as Word Document",
                        data=f,
                        file_name="extracted_text.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

if __name__ == "__main__":
    main()
